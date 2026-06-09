"""Extract plain text from uploaded knowledge-base files."""
from __future__ import annotations

import io
from pathlib import Path


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt", ".xlsx"}


class UnsupportedFileType(Exception):
    pass


class ParseError(Exception):
    pass


def detect_file_type(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "docx"
    if ext == ".md":
        return "md"
    if ext == ".txt":
        return "txt"
    if ext == ".xlsx":
        return "xlsx"
    raise UnsupportedFileType(f"Unsupported file type: {ext or '<no extension>'}")


def extract_text(data: bytes, filename: str) -> str:
    """Return the plain-text content of `data` based on the filename extension."""
    file_type = detect_file_type(filename)
    try:
        if file_type == "pdf":
            return _extract_pdf(data)
        if file_type == "docx":
            return _extract_docx(data)
        if file_type in ("md", "txt"):
            return _extract_plain(data)
        if file_type == "xlsx":
            return _extract_xlsx(data)
    except (UnsupportedFileType, ParseError):
        raise
    except Exception as exc:
        raise ParseError(f"Failed to parse {file_type}: {exc}") from exc
    raise UnsupportedFileType(file_type)


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text.strip())
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_plain(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("utf-8", errors="ignore")


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def chunk_text(text: str, *, chunk_chars: int = 1500, overlap: int = 200) -> list[str]:
    """Naive character-window chunking with overlap. Good enough for
    paragraphs-with-headings docs the merchant uploads as 'about us' material."""
    text = text.strip()
    if not text:
        return []
    chunks: list[str] = []
    start = 0
    n = len(text)
    while start < n:
        end = min(start + chunk_chars, n)
        # Try to break at a paragraph boundary near the end.
        if end < n:
            break_at = text.rfind("\n\n", start + chunk_chars // 2, end)
            if break_at == -1:
                break_at = text.rfind("\n", start + chunk_chars // 2, end)
            if break_at != -1 and break_at > start + 200:
                end = break_at
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end == n:
            break
        start = max(end - overlap, start + 1)
    return chunks
